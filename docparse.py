import PyPDF2
from docx import Document
from openpyxl import load_workbook
import pytesseract
from PIL import Image
import io
import re
import os
import logging
import sys
import olefile
from email import policy
from email.parser import BytesParser
import argparse
import csv
import json
import pandas as pd
import pypff
from bs4 import BeautifulSoup
import validator
import subprocess
import zipfile
import rarfile
import datetime
import signal
import time

def extract_text_from_pdf(pdf_path):

    try:
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)

            text = ""
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()

                    # Проверяем, был ли извлечен текст
                    if page_text:
                        text += page_text
                except BaseException:
                    text = ''
                    pass

            # Проверяем, есть ли извлеченный текст
            if text.strip():  # Если текст не пустой
                return clean_text(text)
            else:  # Если текст пустой, вызываем функцию для извлечения текста из изображений
                print("Текст не найден, пытаемся извлечь текст из изображений.")
                text = clean_text(extract_text_from_pdf_images(pdf_path))
                return (text)
    except Exception as e:
        print(f"Ошибка при чтении PDF как текста: {e}")
        try:
            return clean_text(extract_text_from_pdf_images(pdf_path))
        except Exception as e:
            print(f"Ошибка при извлечении текста из изображений: {e}")
            return None


def extract_text_from_pdf_images(pdf_path):

    try:
        from pdf2image import convert_from_path
        images = convert_from_path(pdf_path)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image, lang='rus+eng')
        return text
    except Exception as e:
        print(f"Ошибка при конвертации PDF в изображения: {e}")
        return None


def extract_text_from_doc(file_path):

    if not os.path.exists("temp"):
        os.makedirs("temp")

    command = [
        "libreoffice",
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        "temp",
        file_path,
    ]

    try:
        subprocess.run(command, check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error converting file: {e}")
        return ""  # Или другое действие при ошибке

    output_docx_path = os.path.join(
        "temp", os.path.splitext(
            os.path.basename(file_path))[0] + ".docx")

    try:
        doc = Document(output_docx_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
    except BaseException:
        text = ""

    os.remove(output_docx_path)

    return text


def extract_text_from_word(file_path: str) -> str:
    """Извлекает текст из .docx файла."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == '.docx':
        # Обработка .docx файлов
        try:
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        except BaseException:
            return None

    elif ext == '.doc':
        # Обработка .doc файлов
        try:
            return extract_text_from_doc(file_path)
        except BaseException:
            return None

    elif ext == '.txt':
        # Обработка .txt файлов
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except BaseException:
            return None

def extract_text_from_excel(file_path: str) -> str:
    _, file_extension = os.path.splitext(file)
    if file_extension.lower() == '.xls':
        if not os.path.exists("temp"):
            os.makedirs("temp")

        command = [
            "libreoffice",
            "--headless",
            "--convert-to",
            "xlsx",
            "--outdir",
            "temp",
            file_path,
        ]

        try:
            subprocess.run(command, check=True)
        except subprocess.CalledProcessError as e:
            print(f"Error converting file: {e}")
            return ""  # Или другое действие при ошибке
        file_path = os.path.join(
            "temp", os.path.splitext(
                os.path.basename(file_path))[0] + ".xlsx")
        print("converting to " + file_path)

    try:
        print("working on xlsx")
        wb = load_workbook(filename=file_path)
        sheet = wb.active
        text = []
        for row in sheet.iter_rows(values_only=True):
            row_text = " ".join(
                [str(cell) if cell is not None else "" for cell in row])
            text.append(row_text)

        return ' '.join(text)
    except BaseException:
        return ''


def find_files_by_extensions(directory, extensions):
    found_files = {"word": [],
                   "excel": [],
                   "pdf": [],
                   "email": [],
                   "archive": []}

    for root, _, files in os.walk(directory):
        for file in files:
            filename, extension = os.path.splitext(file)
            for key in extensions:
                if extension[1:].lower() in [ext.lower()
                                             for ext in extensions[key]]:
                    found_files[key].append(os.path.join(root, file))
    return found_files


def clean_text(text):
    # Убираем знаки табуляции
    text = text.replace('\t', ' ').replace('\n', " ")
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    return text


def sensitive_data_finder(text):
    email_pattern = r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
    phone_pattern = re.compile(
        r'(?:(?:\+7|7|8)[\s-]*)?(?:\(?(\d{3})\)?[\s-]*)?(\d{3})[\s-]*(\d{2})[\s-]*(\d{2})|(\d{10})')
    company_name_pattern = re.compile(
        r'((ООО|ИП|АО|ПАО|ОАО|МОУ|МБОУ|ЗАО|НПАО)( ?)(\"|\«| )[а-яА-Я0-9-_]+(\"|\»| ))',
        re.IGNORECASE)
    ul_inn_pattern = r'[^0-9]\d{10}[^0-9]'
    fl_inn_pattern = r'[^0-9]\d{12}[^0-9]'
    kpp_pattern = r'[^0-9]\d{9}[^0-9]'
    bik_pattern = r'[^0-9]04\d{7}[^0-9]'
    snils_pattern = r'\d{3}-\d{3}-\d{3} \d{2}'
    Full_FIO_pattern = r'[А-Я][а-я]+ [А-Я][а-я]+ [А-Я][а-я]+'
    Abr_FIO_patterns = r'([А-Я](\.|\. | )[А-Я](\.|\. | )[А-Я][а-я]+)'

    emails = re.findall(email_pattern, text)
    matches = phone_pattern.findall(text)
    company_names = re.findall(company_name_pattern, text)
    raw_ul_inn = re.findall(ul_inn_pattern, text)
    raw_fl_inn = re.findall(fl_inn_pattern, text)
    kpp = re.findall(kpp_pattern, text)
    bik = re.findall(bik_pattern, text)
    snilses = re.findall(snils_pattern, text)
    Full_FIOS = re.findall(Full_FIO_pattern, text)
    Abr_FIOS = re.findall(Abr_FIO_patterns, text)
    password_sentenses = extract_password_sentences(text)

    matches = phone_pattern.findall(text)  # Ищем все совпадения

    numbers = []
    for match in matches:
        if match[0]:  # Если найден номер с кодом
            number = ''.join(match)
            if len(number) != 10 or (number[0] not in ["9", "8", "4"]):
                continue
            numbers.append(number)
        elif match[4]:  # Если найден номер без кода (10 цифр)
            number = ''.join(match)
            if len(number) < 10 or len(number) > 12:
                continue
            numbers.append(number)

    company_names = [company_name[0].strip()
                     for company_name in company_names if company_name[0] != '']
    fl_inn = []
    for inn in raw_fl_inn:
        inn = ''.join(filter(str.isdigit, inn))
        if validator.is_valid(inn) and len(inn) != 10:
            fl_inn.append(inn)

    ul_inn = []

    for inn in raw_ul_inn:
        inn = ''.join(filter(str.isdigit, inn))
        if validator.is_valid(inn) and len(inn) != 12:
            ul_inn.append(inn)

    kpp = [inn[1:-1].strip()
           for inn in kpp if inn != '' and not inn[1:-1].startswith('04')]
    bik = [inn[1:-1].strip() for inn in bik if inn != '']
    snilses = [snils.strip() for snils in snilses if snils != '']
    Full_FIOS = [FIO.strip() for FIO in Full_FIOS if FIO != '']
    Abr_FIOS = [FIO[0].strip() for FIO in Abr_FIOS if FIO != '']

    sensitive_data = {
        'phones': numbers,
        'emails': list(set(emails)),
        'company_names': list(set(company_names)),
        'ul_inn': ul_inn,
        'fl_inn': fl_inn,
        'kpp': list(set(kpp)),
        'bik': list(set(bik)),
        'snilses': list(set(snilses)),
        'Full_FIOS': list(set(Full_FIOS)),
        'Abr_FIOS': list(set(Abr_FIOS)),
        'password_sentenses': list(set(password_sentenses))
    }
    return sensitive_data


def extract_password_sentences(text):
    pattern = r'(?<![A-Z])([A-Z][^.!?]*?\b(пароль|password|пароли|passwords|п|пасс|pass)\b[^.!?]*?[.!?])'
    matches = re.findall(pattern, text, re.IGNORECASE)

    sentences = [match[0].strip() for match in matches]

    return sentences


def extract_email_address(header):
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, header)
    return match.group(0) if match else None


def eml_parse(file_path):
    # Открываем файл .eml
    with open(file_path, 'rb') as file:
        # Парсим содержимое файла
        msg = BytesParser(policy=policy.default).parse(file)

    # Извлекаем адрес отправителя и получателей
    sender = msg['From']
    recipients = msg['To']

    # Функция для извлечения адресов из строки
    def extract_addresses(field):
        addresses = re.findall(r'<(.+?)>|([\w\.-]+@[\w\.-]+)', field)
        return [addr[0] or addr[1] for addr in addresses]

    # Извлекаем только адреса
    try:
        sender_address = extract_addresses(sender)[0]  # Предполагаем, что sender всегда один
        recipient_addresses = extract_addresses(recipients)
    except:
        return None

    # Формируем список JSON-объектов
    result = []
    for recipient in recipient_addresses:
        result.append({
            'to': recipient,
            'from': sender_address
        })

    return result

def txt_email_parse(file):
    try:
        with open(file, 'r', encoding='utf-8') as file:
            content = file.read()

        # Регулярки для поиска адресов
        to_pattern = r'To:\s*([\w._%+-]+@[\w.-]+\.[a-zA-Z]{2,}(?:,\s*[\w._%+-]+@[\w.-]+\.[a-zA-Z]{2,})*)'
        from_pattern = r'From:\s*([\w._%+-]+@[\w.-]+\.[a-zA-Z]{2,})'
        to_matches = re.findall(to_pattern, content)
        from_matches = re.findall(from_pattern, content)

        # Обработка найденных адресов
        to_addresses = []
        from_addresses = []

        for match in to_matches:
            # Извлечение адресов из строки
            addresses = [
                email.strip() for email in re.split(
                    r',\s*|\s*;\s*',
                    match) if email]
            to_addresses.extend(addresses)

        for match in from_matches:
            addresses = [
                email.strip() for email in re.split(
                    r',\s*|\s*;\s*',
                    match) if email]
            from_addresses.extend(addresses)
    except BaseException:
        return None
    result = []
    for to_address in to_addresses:
        result.append({
            'to': to_address,
            'from': from_addresses[0]
        })

    return result

def split_and_deduplicate_domains_old(json_data_list):

    unique_data = set()

    for data in json_data_list:
        if isinstance(data['to'], list):
            for to_domain in data['to']:
                new_data = {'from': data['from'], 'to': to_domain}
                unique_data.add(tuple(sorted(new_data.items())))
        else:
            unique_data.add(tuple(sorted(data.items())))

    result = []
    for from_to_pair in unique_data:
        result.append(dict(from_to_pair))
        print(from_to_pair)

    return result


def split_and_deduplicate_domains_bad(data):
    unique_domains = {}
    
    for entry in data:
        from_domain = entry['from']
        to_domain = entry['to']
        
        # Проверяем наличие ключа 'unixtime'
        if 'unixtime' in entry:
            if from_domain not in unique_domains or entry['unixtime'] > unique_domains[from_domain]['unixtime']:
                unique_domains[from_domain] = entry
            if to_domain not in unique_domains or entry['unixtime'] > unique_domains[to_domain]['unixtime']:
                unique_domains[to_domain] = entry

    # Возвращаем только уникальные значения
    return list(unique_domains.values())

def split_and_deduplicate_domains(json_list):
    unique_pairs = {}
    
    for item in json_list:
        key = (item['from'], item['to'])
        # Проверяем наличие ключа 'unixtime' и обновляем, если он больше
        if key not in unique_pairs:
            unique_pairs[key] = item
        else:
            existing_item = unique_pairs[key]
            if 'unixtime' in item and ('unixtime' not in existing_item or item['unixtime'] > existing_item['unixtime']):
                unique_pairs[key] = item

    return list(unique_pairs.values())



def categorize_domains(json_data_list):

    public_email_services = {
        'mail.ru',
        'gmail.com',
        'yandex.ru',
        'rambler.ru',
        'inbox.ru',
        'bk.ru',
        'ya.ru',
        'list.ru',
        "yahoo.com"}
    public_public = []
    public_private = []
    private_private = []

    for data in json_data_list:
        from_domain = data['from']
        to_domain = data['to']

        is_from_public = from_domain in public_email_services
        is_to_public = to_domain in public_email_services

        if is_from_public and is_to_public:
            public_public.append(data)
        elif is_from_public or is_to_public:
            public_private.append(data)
        else:
            private_private.append(data)

    domains = {'public_public': public_public,
               'public_private': public_private,
               'private_private': private_private
               }

    return domains


def extract_from_emails(files):
    pairs = []
    def extract_domains(emails_list):
        result_list = []
        for email_dict in emails_list:
            new_dict = {}
            valid_email_found = False
            for key, value in email_dict.items():
                if key == "unixtime":
                    new_dict['unixtime'] = value
                    continue
                match = re.search(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})", value)
                if match:
                    email = match.group(1)
                    domain_match = re.search(r"@(.*)", email)  
                    if domain_match:
                        domain = domain_match.group(1)
                        new_dict[key] = domain.lower()
                        valid_email_found = True
                    else:
                        valid_email_found = False 
                else:
                    valid_email_found = False
            if valid_email_found:
                result_list.append(new_dict)

        return result_list

    for file in files:
        _, file_extension = os.path.splitext(file)
        if file_extension.lower() == '.txt':
            txt_parse_pairs = txt_email_parse(file) #patch for None return by Except in txt_email_parse
            if txt_parse_pairs != None:              
                pairs.extend(txt_parse_pairs)

        if file_extension.lower() == '.eml':
            eml_parse_pairs = eml_parse(file)
            if eml_parse_pairs != None:
                pairs.extend(eml_parse_pairs)

        if file_extension.lower() == '.pst':
            pst_parse_pairs = pst_parse(file)
            if pst_parse_pairs != None:
                pairs.extend(pst_parse_pairs)

        else:
            continue

    pairs = extract_domains(pairs)

    clear_list = split_and_deduplicate_domains(pairs)
    clear_domains = categorize_domains(clear_list)
    return clear_domains


def export_json_to_csv(mode, json_data, csv_file_path, write_mode='w'):
    if mode == 'sensitive_data':
        # Получаем заголовки из первого элемента JSON
        headers = ['file'] + list(next(iter(json_data.values())).keys())

        with open(csv_file_path, write_mode, newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file)

            # Записываем заголовки
            writer.writerow(headers)

            # Записываем данные
            for file_path, data in json_data.items():
                row = [file_path] + [data[key] for key in headers[1:]]
                writer.writerow(row)
    if mode == 'emails':
        rows = []
        for key in json_data:
            for entry in json_data[key]:
                if 'unixtime' in entry:
                    row = {'type': key, 'from': entry['from'], 'to': entry['to'], 'unixtime':entry['unixtime']}
                else:
                    row = {'type': key, 'from': entry['from'], 'to': entry['to']}
                rows.append(row)
        df = pd.DataFrame(rows)
        df.to_csv(csv_file_path, index=False)


def pst_parse(pst_file_path):
    def extract_addresses(text):
        from_match = re.search(r"From: ([^<]+<([^>]*)>)", text)
        to_match = re.findall(r"To: ([^<]+<([^>]*)>)", text)
        date_string = re.search(r'Sent:\s*(.*?\d{1,2}:\d{2}\s*[AP]M)', text).group(1)
        print(date_string)
        try:
            date_obj = datetime.datetime.strptime(date_string, '%A, %B %d, %Y %I:%M %p')
        except Exception as e:
            print(f"Ошибка при конвертации даты {e}")

        unixtime = int(date_obj.timestamp())

        addresses = []

        if from_match:
            from_address = from_match.group(2)
            if 'unixtime' != None:
                for to_address in [match[1] for match in to_match]:
                    addresses.append({
                        "from": from_address,
                        "to": to_address,
                        "unixtime": unixtime
                    })

            else:
                for to_address in [match[1] for match in to_match]:
                    addresses.append({
                        "from": from_address,
                        "to": to_address,
                    })

        return addresses

    addresses = []
    pst_file = pypff.file()
    pst_file.open(pst_file_path)

    root = pst_file.get_root_folder()

    for folder in root.sub_folders:
        for sub in folder.sub_folders:
            for message in sub.sub_messages:
                try:
                    headers = message.transport_headers
                    raw_body = message.get_html_body()
                except Exception as e:
                    print(f"Ошибка при получении тела письма {e}")
                    raw_body = None
                    continue
                if raw_body is not None:
                    # print("stststs" + str(headers))
                    # print(raw_body)
                    try:
                        raw_body = raw_body.decode('Windows-1251')
                        # print(raw_body)
                    except BaseException:
                        raw_body = raw_body.decode('utf-8')
                        # print(raw_body)
                    soup = BeautifulSoup(raw_body, "lxml")
                    plain_text = soup.get_text()
                    try:
                        addresses.extend(extract_addresses(plain_text))
                    except:
                        continue

    return addresses


def process_archives(archives, extensions, path):
    temp_archives = path + "/temp/archives"
    os.makedirs(temp_archives, exist_ok=True)

    result = {}
    for category, extensions_list in extensions.items():
        result[category] = []
        for archive_path in archives['archive']:
            try:
                if archive_path.lower().endswith(".zip"):
                    with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                        for file_info in zip_ref.infolist():
                            if file_info.filename.lower().endswith(tuple(extensions_list)):
                                file_path = os.path.join(
                                    temp_archives, file_info.filename)
                                zip_ref.extract(file_info, temp_archives)
                                result[category].append(file_path)
                elif archive_path.lower().endswith(".rar"):
                    with rarfile.RarFile(archive_path, 'r') as rar_ref:
                        for file_info in rar_ref.infolist():
                            if file_info.filename.lower().endswith(tuple(extensions_list)):
                                file_path = os.path.join(
                                    temp_archives, file_info.filename)
                                rar_ref.extract(file_info, temp_archives)
                                result[category].append(file_path)
            except Exception as e:
                print(f"Ошибка при обработке архива {archive_path}: {e}")

    return result

def handler(signum, frame):
    raise TimeoutError("Time's up!")

def run_with_timeout(func, timeout_seconds):
    signal.signal(signal.SIGALRM, handler)
    signal.alarm(timeout_seconds)
    try:
        return func()  # Возвращаем результат функции
    except TimeoutError:
        print("Operation timed out!")
        return None  # Возвращаем None в случае таймаута
    finally:
        signal.alarm(0)

if __name__ == "__main__":
    extensions = {"word": ["txt", "docx", "doc"],
                  "excel": ["xlsx", "xls"],
                  "pdf": ["pdf"],
                  "email": ["pst", "txt", "eml"],
                  "archive": ["rar", "zip"]}

    parser = argparse.ArgumentParser(
        description="Extract emails or sensitive data from a directory.")
    parser.add_argument(
        '--input_path',
        type=str,
        required=True,
        help='Path to the input directory')
    parser.add_argument('--mode', type=str, choices=['emails', 'sensitive_data'], required=True,
                        help='Mode of operation: "emails" (search for From-To emails and domains in addresses) or "sensitive_data" (search for phones, names etc)')
    args = parser.parse_args()
    path = args.input_path
#   archives = find_files_by_extensions(
#   path, {'archive': extensions['archive']})
#   process_archives(archives, extensions, path)
    files = find_files_by_extensions(path, extensions)
    results = {}
    problem_files = []
    csv_file_path = 'output1.csv'

    if args.mode == "emails":
        results = extract_from_emails(files['email'])

    if args.mode == "sensitive_data":
        number_of_files = len(files['word'])
        for file in files['word']:
            print("File num " + str(files['word'].index(file)) + "/" + str(number_of_files))
            text = run_with_timeout(lambda: extract_text_from_word(file), 1)
            if text != None:
                sensitive_data = sensitive_data_finder(text)
                if sensitive_data != None:
                    results[file] = sensitive_data
            else:
                problem_files.append(file)
                continue
        export_json_to_csv(args.mode, results, csv_file_path)
        results = {}

        for file in files['excel']:
            print("File num " + str(files['excel'].index(file)) + "/" + str(number_of_files))
            text = run_with_timeout(lambda: extract_text_from_pdf(file), 1)
            if text != None:
                sensitive_data = run_with_timeout(lambda: sensitive_data_finder(text), 1)
                if sensitive_data != None:
                    results[file] = sensitive_data
            else:
                problem_files.append(file)
                continue
        export_json_to_csv(args.mode, results, csv_file_path, "a")
        results = {}

        for file in files['pdf']:
            print("File num " + str(files['pdf'].index(file)) + "/" + str(number_of_files))
            text = run_with_timeout(lambda: extract_text_from_pdf(file), 1)
            if text != None:
                sensitive_data = run_with_timeout(lambda: sensitive_data_finder(text), 1)
                if sensitive_data != None:
                    results[file] = sensitive_data
            else:
                problem_files.append(file)
                continue
        export_json_to_csv(args.mode, results, csv_file_path, 'a')
        print(problem_files)



    csv_file_path = 'output1.csv'
    export_json_to_csv(args.mode, results, csv_file_path)
    with open('output1.json', 'w', encoding='utf-8') as json_file:
        json.dump(results, json_file, ensure_ascii=False, indent=4)
